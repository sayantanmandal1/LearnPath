"""
Resume processing service for handling file uploads, text extraction, and AI-powered parsing
"""
import os
import asyncio
import logging
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path
import tempfile
import mimetypes

import fitz  # PyMuPDF
from docx import Document
import httpx
from fastapi import UploadFile, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update

from app.models.resume import ResumeData, ProcessingStatus
from app.schemas.resume import (
    ParsedResumeData, ContactInfo, WorkExperience, Education, 
    Certification, SkillCategory, ResumeValidationResult, 
    ResumeValidationError, SupportedFileType
)
from app.core.database import get_db
from app.core.exceptions import ValidationError, ProcessingError

logger = logging.getLogger(__name__)


class ResumeProcessingService:
    """Service for processing resume uploads and extracting structured data"""
    
    def __init__(self):
        self.upload_dir = Path("uploads/resumes")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.supported_types = {
            SupportedFileType.PDF.value,
            SupportedFileType.DOC.value,
            SupportedFileType.DOCX.value
        }
        
        # Gemini API configuration (placeholder - would be configured via environment)
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        self.gemini_api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent"
    
    async def upload_resume(self, file: UploadFile, user_id: str, db: AsyncSession) -> ResumeData:
        """
        Handle resume file upload with validation
        
        Args:
            file: Uploaded file object
            user_id: User ID for the resume
            db: Database session
            
        Returns:
            ResumeData: Created resume record
            
        Raises:
            ValidationError: If file validation fails
            ProcessingError: If file storage fails
        """
        try:
            # Validate file
            await self._validate_file(file)
            
            # Generate unique filename
            file_extension = Path(file.filename).suffix.lower()
            unique_filename = f"{user_id}_{datetime.utcnow().timestamp()}{file_extension}"
            file_path = self.upload_dir / unique_filename
            
            # Save file to disk
            file_content = await file.read()
            with open(file_path, "wb") as f:
                f.write(file_content)
            
            # Create database record
            resume_data = ResumeData(
                user_id=user_id,
                original_filename=file.filename,
                file_path=str(file_path),
                file_size=len(file_content),
                file_type=file.content_type,
                processing_status=ProcessingStatus.PENDING
            )
            
            db.add(resume_data)
            await db.commit()
            await db.refresh(resume_data)
            
            logger.info(f"Resume uploaded successfully: {resume_data.id}")
            return resume_data
            
        except Exception as e:
            logger.error(f"Resume upload failed: {str(e)}")
            raise ProcessingError(f"Failed to upload resume: {str(e)}")
    
    async def _validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file format and size
        
        Args:
            file: Uploaded file to validate
            
        Raises:
            ValidationError: If validation fails
        """
        # Check file size
        file_content = await file.read()
        await file.seek(0)  # Reset file pointer
        
        if len(file_content) > self.max_file_size:
            raise ValidationError(f"File size exceeds maximum limit of {self.max_file_size / (1024*1024):.1f}MB")
        
        if len(file_content) == 0:
            raise ValidationError("File is empty")
        
        # Check file type
        if file.content_type not in self.supported_types:
            # Try to detect MIME type from filename
            detected_type, _ = mimetypes.guess_type(file.filename)
            if detected_type not in self.supported_types:
                raise ValidationError(
                    f"Unsupported file type: {file.content_type}. "
                    f"Supported types: PDF, DOC, DOCX"
                )
            file.content_type = detected_type
        
        # Basic file header validation
        if file.content_type == SupportedFileType.PDF.value and not file_content.startswith(b'%PDF'):
            raise ValidationError("Invalid PDF file format")
    
    async def process_resume(self, resume_id: str, db: AsyncSession) -> ResumeData:
        """
        Process resume by extracting text and parsing with Gemini API
        
        Args:
            resume_id: ID of resume to process
            db: Database session
            
        Returns:
            ResumeData: Updated resume record with processing results
        """
        try:
            # Get resume record
            result = await db.execute(select(ResumeData).where(ResumeData.id == resume_id))
            resume_data = result.scalar_one_or_none()
            
            if not resume_data:
                raise ProcessingError(f"Resume not found: {resume_id}")
            
            # Update status to processing
            await self._update_processing_status(
                db, resume_id, ProcessingStatus.PROCESSING, 
                processing_started_at=datetime.utcnow()
            )
            
            # Extract text from file
            extracted_text, confidence = await self._extract_text_from_file(
                resume_data.file_path, resume_data.file_type
            )
            
            # Parse with Gemini API
            parsed_data = await self._parse_with_gemini(extracted_text)
            
            # Validate and normalize parsed data
            validation_result = await self._validate_parsed_data(parsed_data)
            
            # Update database with results
            await self._update_resume_with_results(
                db, resume_id, extracted_text, confidence, 
                parsed_data, validation_result
            )
            
            # Get updated record
            result = await db.execute(select(ResumeData).where(ResumeData.id == resume_id))
            updated_resume = result.scalar_one()
            
            logger.info(f"Resume processed successfully: {resume_id}")
            return updated_resume
            
        except Exception as e:
            logger.error(f"Resume processing failed for {resume_id}: {str(e)}")
            await self._update_processing_status(
                db, resume_id, ProcessingStatus.FAILED,
                error_message=str(e),
                processing_completed_at=datetime.utcnow()
            )
            raise ProcessingError(f"Failed to process resume: {str(e)}")
    
    async def _extract_text_from_file(self, file_path: str, file_type: str) -> Tuple[str, float]:
        """
        Extract text from resume file based on file type
        
        Args:
            file_path: Path to the resume file
            file_type: MIME type of the file
            
        Returns:
            Tuple[str, float]: Extracted text and confidence score
        """
        try:
            if file_type == SupportedFileType.PDF.value:
                return await self._extract_from_pdf(file_path)
            elif file_type in [SupportedFileType.DOC.value, SupportedFileType.DOCX.value]:
                return await self._extract_from_docx(file_path)
            else:
                raise ProcessingError(f"Unsupported file type for extraction: {file_type}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise ProcessingError(f"Failed to extract text: {str(e)}")
    
    async def _extract_from_pdf(self, file_path: str) -> Tuple[str, float]:
        """Extract text from PDF file"""
        try:
            doc = fitz.open(file_path)
            text_blocks = []
            total_chars = 0
            extracted_chars = 0
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_blocks.append(text)
                    extracted_chars += len(text)
                total_chars += len(page.get_text("rawdict")["blocks"])
            
            doc.close()
            
            extracted_text = "\n\n".join(text_blocks)
            confidence = min(0.95, extracted_chars / max(total_chars, 1)) if total_chars > 0 else 0.0
            
            return extracted_text, confidence
            
        except Exception as e:
            raise ProcessingError(f"PDF extraction failed: {str(e)}")
    
    async def _extract_from_docx(self, file_path: str) -> Tuple[str, float]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(paragraphs)
            confidence = 0.9 if extracted_text.strip() else 0.0
            
            return extracted_text, confidence
            
        except Exception as e:
            raise ProcessingError(f"DOCX extraction failed: {str(e)}")
    
    async def _parse_with_gemini(self, text: str) -> ParsedResumeData:
        """
        Parse extracted text using Gemini API with enhanced error handling
        
        Args:
            text: Extracted resume text
            
        Returns:
            ParsedResumeData: Structured resume data
        """
        from app.core.error_handling_decorators import with_gemini_error_handling
        from app.core.exceptions import MLModelError
        
        if not self.gemini_api_key:
            logger.warning("Gemini API key not configured, using fallback parsing")
            return await self._fallback_parsing(text)
        
        try:
            prompt = self._create_parsing_prompt(text)
            
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.post(
                    f"{self.gemini_api_url}?key={self.gemini_api_key}",
                    json={
                        "contents": [{
                            "parts": [{"text": prompt}]
                        }],
                        "generationConfig": {
                            "temperature": 0.1,
                            "maxOutputTokens": 2048
                        }
                    }
                )
                
                # Enhanced error handling for different response codes
                if response.status_code == 401:
                    logger.error("Gemini API authentication failed")
                    return await self._fallback_parsing(text)
                elif response.status_code == 403:
                    logger.error("Gemini API access forbidden")
                    return await self._fallback_parsing(text)
                elif response.status_code == 429:
                    logger.warning("Gemini API rate limit exceeded, using fallback")
                    return await self._fallback_parsing(text)
                elif response.status_code == 503:
                    logger.warning("Gemini API service unavailable, using fallback")
                    return await self._fallback_parsing(text)
                elif response.status_code != 200:
                    logger.error(f"Gemini API error: {response.status_code} - {response.text}")
                    return await self._fallback_parsing(text)
                
                result = response.json()
                
                # Check for API errors in response
                if "error" in result:
                    logger.error(f"Gemini API returned error: {result['error']}")
                    return await self._fallback_parsing(text)
                
                if "candidates" not in result or not result["candidates"]:
                    logger.warning("Gemini API returned no candidates, using fallback")
                    return await self._fallback_parsing(text)
                
                candidate = result["candidates"][0]
                
                # Check for content filtering
                if candidate.get("finishReason") == "SAFETY":
                    logger.warning("Gemini API filtered content for safety, using fallback")
                    return await self._fallback_parsing(text)
                
                if "content" not in candidate or "parts" not in candidate["content"]:
                    logger.warning("Gemini API returned malformed response, using fallback")
                    return await self._fallback_parsing(text)
                
                parsed_content = candidate["content"]["parts"][0]["text"]
                
                return await self._parse_gemini_response(parsed_content)
                
        except httpx.TimeoutException:
            logger.warning("Gemini API request timed out, using fallback parsing")
            return await self._fallback_parsing(text)
        except httpx.ConnectError:
            logger.warning("Failed to connect to Gemini API, using fallback parsing")
            return await self._fallback_parsing(text)
        except Exception as e:
            logger.error(f"Gemini API parsing failed: {str(e)}")
            return await self._fallback_parsing(text)
    
    def _create_parsing_prompt(self, text: str) -> str:
        """Create structured prompt for Gemini API"""
        return f"""
        Please parse the following resume text and extract structured information in JSON format.
        
        Extract the following sections:
        1. Contact Information (name, email, phone, location, linkedin, github, portfolio)
        2. Professional Summary
        3. Work Experience (company, position, dates, description, technologies, achievements)
        4. Education (institution, degree, field, dates, gpa, achievements)
        5. Skills (categorized by type with proficiency levels)
        6. Certifications (name, issuer, dates, credential_id)
        7. Projects (name, description, technologies, url)
        8. Languages
        9. Awards and Achievements
        
        Resume Text:
        {text}
        
        Please respond with valid JSON only, no additional text.
        """
    
    async def _parse_gemini_response(self, response_text: str) -> ParsedResumeData:
        """Parse Gemini API response into structured data"""
        try:
            import json
            
            # Clean response text (remove markdown formatting if present)
            clean_text = response_text.strip()
            if clean_text.startswith("```json"):
                clean_text = clean_text[7:]
            if clean_text.endswith("```"):
                clean_text = clean_text[:-3]
            
            parsed_json = json.loads(clean_text)
            
            # Convert to Pydantic models
            return ParsedResumeData(
                contact_info=ContactInfo(**parsed_json.get("contact_info", {})) if parsed_json.get("contact_info") else None,
                summary=parsed_json.get("summary"),
                work_experience=[WorkExperience(**exp) for exp in parsed_json.get("work_experience", [])],
                education=[Education(**edu) for edu in parsed_json.get("education", [])],
                skills=[SkillCategory(**skill) for skill in parsed_json.get("skills", [])],
                certifications=[Certification(**cert) for cert in parsed_json.get("certifications", [])],
                projects=parsed_json.get("projects", []),
                languages=parsed_json.get("languages", []),
                awards=parsed_json.get("awards", [])
            )
            
        except Exception as e:
            logger.error(f"Failed to parse Gemini response: {str(e)}")
            return await self._fallback_parsing(response_text)
    
    async def _fallback_parsing(self, text: str) -> ParsedResumeData:
        """Fallback parsing when Gemini API is unavailable"""
        # Basic text analysis for fallback
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Simple email extraction
        import re
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        # Simple phone extraction
        phone_pattern = r'[\+]?[1-9]?[0-9]{7,15}'
        phones = re.findall(phone_pattern, text)
        
        contact_info = ContactInfo(
            email=emails[0] if emails else None,
            phone=phones[0] if phones else None
        )
        
        return ParsedResumeData(
            contact_info=contact_info,
            summary="Automatic parsing unavailable. Please review and update manually."
        )
    
    async def _validate_parsed_data(self, parsed_data: ParsedResumeData) -> ResumeValidationResult:
        """Validate and score parsed resume data"""
        errors = []
        warnings = []
        score_components = []
        
        # Validate contact information
        if parsed_data.contact_info:
            if not parsed_data.contact_info.email:
                warnings.append(ResumeValidationError(
                    field="contact_info.email",
                    message="Email address not found"
                ))
            else:
                score_components.append(0.2)
                
            if not parsed_data.contact_info.phone:
                warnings.append(ResumeValidationError(
                    field="contact_info.phone", 
                    message="Phone number not found"
                ))
            else:
                score_components.append(0.1)
        else:
            errors.append(ResumeValidationError(
                field="contact_info",
                message="No contact information found"
            ))
        
        # Validate work experience
        if parsed_data.work_experience:
            score_components.append(0.3)
            for i, exp in enumerate(parsed_data.work_experience):
                if not exp.company:
                    warnings.append(ResumeValidationError(
                        field=f"work_experience[{i}].company",
                        message="Company name missing"
                    ))
        else:
            warnings.append(ResumeValidationError(
                field="work_experience",
                message="No work experience found"
            ))
        
        # Validate education
        if parsed_data.education:
            score_components.append(0.2)
        
        # Validate skills
        if parsed_data.skills:
            score_components.append(0.2)
        
        confidence_score = sum(score_components)
        
        return ResumeValidationResult(
            is_valid=len(errors) == 0,
            errors=errors,
            warnings=warnings,
            confidence_score=confidence_score
        )
    
    async def _update_processing_status(
        self, 
        db: AsyncSession, 
        resume_id: str, 
        status: ProcessingStatus,
        error_message: Optional[str] = None,
        processing_started_at: Optional[datetime] = None,
        processing_completed_at: Optional[datetime] = None
    ) -> None:
        """Update resume processing status"""
        update_data = {"processing_status": status, "updated_at": datetime.utcnow()}
        
        if error_message:
            update_data["error_message"] = error_message
        if processing_started_at:
            update_data["processing_started_at"] = processing_started_at
        if processing_completed_at:
            update_data["processing_completed_at"] = processing_completed_at
            
        await db.execute(
            update(ResumeData)
            .where(ResumeData.id == resume_id)
            .values(**update_data)
        )
        await db.commit()
    
    async def _update_resume_with_results(
        self,
        db: AsyncSession,
        resume_id: str,
        extracted_text: str,
        confidence: float,
        parsed_data: ParsedResumeData,
        validation_result: ResumeValidationResult
    ) -> None:
        """Update resume record with processing results"""
        update_data = {
            "processing_status": ProcessingStatus.COMPLETED,
            "extracted_text": extracted_text,
            "extraction_confidence": confidence,
            "parsed_sections": parsed_data.dict(),
            "contact_info": parsed_data.contact_info.dict() if parsed_data.contact_info else None,
            "work_experience": [exp.dict() for exp in parsed_data.work_experience] if parsed_data.work_experience else None,
            "education_data": [edu.dict() for edu in parsed_data.education] if parsed_data.education else None,
            "skills_extracted": [skill.dict() for skill in parsed_data.skills] if parsed_data.skills else None,
            "certifications_data": [cert.dict() for cert in parsed_data.certifications] if parsed_data.certifications else None,
            "processing_completed_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        await db.execute(
            update(ResumeData)
            .where(ResumeData.id == resume_id)
            .values(**update_data)
        )
        await db.commit()
    
    async def get_resume_by_id(self, resume_id: str, db: AsyncSession) -> Optional[ResumeData]:
        """Get resume by ID"""
        result = await db.execute(select(ResumeData).where(ResumeData.id == resume_id))
        return result.scalar_one_or_none()
    
    async def get_user_resumes(self, user_id: str, db: AsyncSession) -> List[ResumeData]:
        """Get all resumes for a user"""
        result = await db.execute(
            select(ResumeData)
            .where(ResumeData.user_id == user_id)
            .order_by(ResumeData.created_at.desc())
        )
        return result.scalars().all()