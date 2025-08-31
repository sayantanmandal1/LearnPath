"""
Comprehensive tests for the NLP processing engine.
"""

import pytest
import asyncio
import tempfile
import os
from pathlib import Path
from unittest.mock import Mock, patch, AsyncMock

import numpy as np
from reportlab.pdfgen import canvas
from docx import Document

from ..nlp_engine import NLPEngine
from ..models import ResumeData, SkillExtraction, SkillCategory


class TestNLPEngine:
    """Test suite for NLP Engine."""
    
    @pytest.fixture
    def nlp_engine(self):
        """Create NLP engine instance for testing."""
        with tempfile.TemporaryDirectory() as temp_dir:
            engine = NLPEngine(model_cache_dir=temp_dir)
            yield engine
            engine.cleanup()
    
    @pytest.fixture
    def sample_resume_text(self):
        """Sample resume text for testing."""
        return """
        John Doe
        Software Engineer
        john.doe@email.com
        (555) 123-4567
        
        EXPERIENCE
        Senior Software Engineer at TechCorp (2020-2023)
        - Developed web applications using Python, Django, and React
        - Worked with PostgreSQL and Redis for data storage
        - Implemented CI/CD pipelines using Docker and Kubernetes
        - Led a team of 5 developers using Agile methodologies
        
        Software Developer at StartupXYZ (2018-2020)
        - Built REST APIs using Node.js and Express
        - Used MongoDB for database management
        - Implemented authentication using JWT tokens
        
        EDUCATION
        Bachelor of Science in Computer Science
        University of Technology (2014-2018)
        
        SKILLS
        Programming Languages: Python, JavaScript, TypeScript, Java
        Frameworks: Django, React, Node.js, Express
        Databases: PostgreSQL, MongoDB, Redis
        Cloud: AWS, Docker, Kubernetes
        Tools: Git, Jenkins, JIRA
        
        CERTIFICATIONS
        AWS Certified Solutions Architect
        Certified Kubernetes Administrator
        """   
 
    @pytest.fixture
    def sample_pdf_file(self):
        """Create a sample PDF file for testing."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
            # Create a simple PDF with sample text
            c = canvas.Canvas(temp_file.name)
            c.drawString(100, 750, "John Doe - Software Engineer")
            c.drawString(100, 730, "Skills: Python, JavaScript, React, Django")
            c.drawString(100, 710, "Experience: 5 years in web development")
            c.save()
            
            yield temp_file.name
            
            # Cleanup
            try:
                os.unlink(temp_file.name)
            except OSError:
                pass
    
    @pytest.fixture
    def sample_docx_file(self):
        """Create a sample DOCX file for testing."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc = Document()
            doc.add_heading('John Doe - Software Engineer', 0)
            doc.add_paragraph('Skills: Python, JavaScript, React, Django')
            doc.add_paragraph('Experience: 5 years in web development')
            doc.add_paragraph('Education: BS Computer Science')
            doc.save(temp_file.name)
            
            yield temp_file.name
            
            # Cleanup
            try:
                os.unlink(temp_file.name)
            except OSError:
                pass

    # Test Resume Parsing
    @pytest.mark.asyncio
    async def test_parse_pdf_resume(self, nlp_engine, sample_pdf_file):
        """Test PDF resume parsing."""
        resume_data = await nlp_engine.parse_resume(sample_pdf_file, 'pdf')
        
        assert isinstance(resume_data, ResumeData)
        assert resume_data.file_path == sample_pdf_file
        assert resume_data.file_type == 'pdf'
        assert len(resume_data.raw_text) > 0
        assert 'John Doe' in resume_data.raw_text
        assert 'Python' in resume_data.raw_text
    
    @pytest.mark.asyncio
    async def test_parse_docx_resume(self, nlp_engine, sample_docx_file):
        """Test DOCX resume parsing."""
        resume_data = await nlp_engine.parse_resume(sample_docx_file, 'docx')
        
        assert isinstance(resume_data, ResumeData)
        assert resume_data.file_path == sample_docx_file
        assert resume_data.file_type == 'docx'
        assert len(resume_data.raw_text) > 0
        assert 'John Doe' in resume_data.raw_text
        assert 'Python' in resume_data.raw_text
    
    @pytest.mark.asyncio
    async def test_parse_resume_auto_detect_type(self, nlp_engine, sample_pdf_file):
        """Test automatic file type detection."""
        resume_data = await nlp_engine.parse_resume(sample_pdf_file)
        
        assert resume_data.file_type == 'pdf'
    
    @pytest.mark.asyncio
    async def test_parse_resume_unsupported_type(self, nlp_engine):
        """Test handling of unsupported file types."""
        with tempfile.NamedTemporaryFile(suffix='.txt') as temp_file:
            with pytest.raises(ValueError, match="Unsupported file type"):
                await nlp_engine.parse_resume(temp_file.name, 'txt')
    
    # Test Text Processing
    @pytest.mark.asyncio
    async def test_process_resume_text(self, nlp_engine, sample_resume_text):
        """Test resume text processing."""
        resume_data = await nlp_engine._process_resume_text(sample_resume_text)
        
        assert isinstance(resume_data, ResumeData)
        assert len(resume_data.skills) > 0
        assert len(resume_data.experience) > 0
        assert len(resume_data.education) > 0
        assert len(resume_data.certifications) > 0
    
    def test_extract_sections(self, nlp_engine, sample_resume_text):
        """Test section extraction from resume text."""
        sections = nlp_engine._extract_sections(sample_resume_text)
        
        assert 'experience' in sections
        assert 'education' in sections
        assert 'skills' in sections
        assert 'certifications' in sections
        
        assert 'TechCorp' in sections['experience']
        assert 'Computer Science' in sections['education']
        assert 'Python' in sections['skills']
        assert 'AWS Certified' in sections['certifications']
    
    # Test Skill Extraction
    @pytest.mark.asyncio
    async def test_extract_skills(self, nlp_engine, sample_resume_text):
        """Test skill extraction from text."""
        skills = await nlp_engine.extract_skills(sample_resume_text)
        
        assert len(skills) > 0
        skill_names = [skill.skill_name for skill in skills]
        
        # Check for expected skills
        expected_skills = ['Python', 'JavaScript', 'React', 'Django', 'PostgreSQL']
        for expected_skill in expected_skills:
            assert any(expected_skill.lower() in skill.lower() for skill in skill_names)
        
        # Check skill properties
        for skill in skills:
            assert isinstance(skill, SkillExtraction)
            assert 0.0 <= skill.confidence_score <= 1.0
            assert skill.source in ['spacy_ner', 'pattern_matching', 'bert_classification']
            assert isinstance(skill.category, SkillCategory)
    
    @pytest.mark.asyncio
    async def test_extract_skills_spacy(self, nlp_engine):
        """Test spaCy-based skill extraction."""
        text = "I have experience with Python programming and React development."
        skills = await nlp_engine._extract_skills_spacy(text)
        
        assert isinstance(skills, list)
        # Note: Results may vary based on spaCy model availability
    
    def test_extract_skills_patterns(self, nlp_engine):
        """Test pattern-based skill extraction."""
        text = "I work with Python, JavaScript, React, Django, AWS, and PostgreSQL."
        skills = nlp_engine._extract_skills_patterns(text)
        
        assert len(skills) > 0
        skill_names = [skill.skill_name for skill in skills]
        
        expected_skills = ['Python', 'JavaScript', 'React', 'Django', 'AWS', 'PostgreSQL']
        for expected_skill in expected_skills:
            assert expected_skill in skill_names
    
    @pytest.mark.asyncio
    async def test_extract_skills_bert(self, nlp_engine):
        """Test BERT-based skill extraction."""
        text = "I have experience with Python programming and machine learning frameworks."
        skills = await nlp_engine._extract_skills_bert(text)
        
        assert isinstance(skills, list)
        # Results may vary based on model availability
    
    def test_is_technical_skill(self, nlp_engine):
        """Test technical skill identification."""
        assert nlp_engine._is_technical_skill("Python")
        assert nlp_engine._is_technical_skill("React.js")
        assert nlp_engine._is_technical_skill("AWS")
        assert nlp_engine._is_technical_skill("PostgreSQL")
        
        assert not nlp_engine._is_technical_skill("the")
        assert not nlp_engine._is_technical_skill("and")
        assert not nlp_engine._is_technical_skill("very")
    
    def test_is_skill_sentence(self, nlp_engine):
        """Test skill sentence identification."""
        assert nlp_engine._is_skill_sentence("I have experience with Python programming")
        assert nlp_engine._is_skill_sentence("Proficient in JavaScript and React")
        assert nlp_engine._is_skill_sentence("Skilled in database management")
        
        assert not nlp_engine._is_skill_sentence("I went to the store yesterday")
        assert not nlp_engine._is_skill_sentence("The weather is nice today")
    
    def test_merge_skills(self, nlp_engine):
        """Test skill merging and deduplication."""
        skills = [
            SkillExtraction(
                skill_name="Python",
                confidence_score=0.8,
                source="pattern",
                evidence=["Python programming"],
                category=SkillCategory.PROGRAMMING_LANGUAGES
            ),
            SkillExtraction(
                skill_name="python",
                confidence_score=0.7,
                source="spacy",
                evidence=["python development"],
                category=SkillCategory.PROGRAMMING_LANGUAGES
            ),
            SkillExtraction(
                skill_name="JavaScript",
                confidence_score=0.9,
                source="pattern",
                evidence=["JavaScript coding"],
                category=SkillCategory.PROGRAMMING_LANGUAGES
            )
        ]
        
        merged = nlp_engine._merge_skills(skills)
        
        assert len(merged) == 2  # Python entries should be merged
        
        # Find the merged Python skill
        python_skill = next(skill for skill in merged if skill.skill_name.lower() == "python")
        assert python_skill.confidence_score == 0.8  # Should take max confidence
        assert len(python_skill.evidence) == 2  # Should combine evidence
        assert "pattern, spacy" in python_skill.source
    
    # Test Experience Extraction
    def test_extract_experience(self, nlp_engine, sample_resume_text):
        """Test work experience extraction."""
        sections = nlp_engine._extract_sections(sample_resume_text)
        experience = nlp_engine._extract_experience(sections.get('experience', ''))
        
        assert len(experience) >= 2  # Should find at least 2 job entries
        
        # Check first experience entry
        first_exp = experience[0]
        assert 'TechCorp' in first_exp['raw_text'] or 'StartupXYZ' in first_exp['raw_text']
        assert first_exp['dates'] is not None or first_exp['first_line'] is not None
    
    def test_parse_experience_entry(self, nlp_engine):
        """Test parsing of individual experience entries."""
        entry_text = """
        Senior Software Engineer at TechCorp (2020-2023)
        - Developed web applications using Python and Django
        - Led a team of 5 developers
        """
        
        parsed = nlp_engine._parse_experience_entry(entry_text.strip())
        
        assert parsed is not None
        assert 'TechCorp' in parsed['first_line']
        assert '2020-2023' in parsed['dates']
        assert 'Python' in parsed['description']
    
    # Test Education Extraction
    def test_extract_education(self, nlp_engine, sample_resume_text):
        """Test education extraction."""
        sections = nlp_engine._extract_sections(sample_resume_text)
        education = nlp_engine._extract_education(sections.get('education', ''))
        
        assert len(education) >= 1
        
        # Check education entry
        first_edu = education[0]
        assert 'Computer Science' in first_edu['degree_institution']
        assert first_edu['graduation_year'] is not None
    
    def test_parse_education_entry(self, nlp_engine):
        """Test parsing of individual education entries."""
        entry_text = """
        Bachelor of Science in Computer Science
        University of Technology (2014-2018)
        """
        
        parsed = nlp_engine._parse_education_entry(entry_text.strip())
        
        assert parsed is not None
        assert 'Computer Science' in parsed['degree_institution']
        assert '2018' in parsed['graduation_year']
    
    # Test Certification Extraction
    def test_extract_certifications(self, nlp_engine, sample_resume_text):
        """Test certification extraction."""
        certifications = nlp_engine._extract_certifications(sample_resume_text)
        
        assert len(certifications) >= 2
        cert_text = ' '.join(certifications).lower()
        assert 'aws' in cert_text
        assert 'kubernetes' in cert_text
    
    # Test Embedding Generation
    @pytest.mark.asyncio
    async def test_generate_embeddings(self, nlp_engine):
        """Test semantic embedding generation."""
        texts = [
            "Python programming language",
            "JavaScript web development",
            "Machine learning with TensorFlow"
        ]
        
        embeddings = await nlp_engine.generate_embeddings(texts)
        
        assert isinstance(embeddings, np.ndarray)
        assert embeddings.shape[0] == len(texts)
        assert embeddings.shape[1] > 0  # Should have embedding dimensions
    
    @pytest.mark.asyncio
    async def test_calculate_similarity(self, nlp_engine):
        """Test semantic similarity calculation."""
        text1 = "Python programming language"
        text2 = "Python software development"
        text3 = "Cooking recipes and food"
        
        # Similar texts should have high similarity
        similarity_high = await nlp_engine.calculate_similarity(text1, text2)
        assert 0.5 <= similarity_high <= 1.0
        
        # Dissimilar texts should have low similarity
        similarity_low = await nlp_engine.calculate_similarity(text1, text3)
        assert 0.0 <= similarity_low <= 0.5
    
    # Test Edge Cases and Error Handling
    @pytest.mark.asyncio
    async def test_empty_text_processing(self, nlp_engine):
        """Test handling of empty text."""
        resume_data = await nlp_engine._process_resume_text("")
        
        assert isinstance(resume_data, ResumeData)
        assert resume_data.raw_text == ""
        assert resume_data.cleaned_text == ""
        assert len(resume_data.skills) == 0
    
    @pytest.mark.asyncio
    async def test_malformed_resume_text(self, nlp_engine):
        """Test handling of malformed resume text."""
        malformed_text = "!@#$%^&*()_+{}|:<>?[]\\;'\",./"
        
        resume_data = await nlp_engine._process_resume_text(malformed_text)
        
        assert isinstance(resume_data, ResumeData)
        # Should handle gracefully without crashing
    
    @pytest.mark.asyncio
    async def test_very_long_text(self, nlp_engine):
        """Test handling of very long text."""
        long_text = "Python programming. " * 1000  # Very long text
        
        skills = await nlp_engine.extract_skills(long_text)
        
        assert isinstance(skills, list)
        # Should handle without memory issues
    
    def test_get_context(self, nlp_engine):
        """Test context extraction around matches."""
        text = "I have extensive experience with Python programming and web development using Django framework."
        start = text.find("Python")
        end = start + len("Python")
        
        context = nlp_engine._get_context(text, start, end, context_size=20)
        
        assert "Python" in context
        assert len(context) <= len(text)
    
    # Test Utility Functions
    def test_extract_noun_phrases(self, nlp_engine):
        """Test noun phrase extraction."""
        text = "I work with machine learning algorithms and deep neural networks."
        
        phrases = nlp_engine._extract_noun_phrases(text)
        
        assert isinstance(phrases, list)
        # Results may vary based on spaCy model
    
    # Test Performance and Accuracy
    @pytest.mark.asyncio
    async def test_skill_extraction_accuracy(self, nlp_engine):
        """Test skill extraction accuracy with known skills."""
        text = """
        I am proficient in Python, JavaScript, React, Django, PostgreSQL, 
        AWS, Docker, Kubernetes, Git, and Jenkins. I have experience with 
        machine learning using TensorFlow and PyTorch.
        """
        
        skills = await nlp_engine.extract_skills(text)
        skill_names = [skill.skill_name.lower() for skill in skills]
        
        # Expected skills that should be found
        expected_skills = [
            'python', 'javascript', 'react', 'django', 'postgresql',
            'aws', 'docker', 'kubernetes', 'git', 'jenkins'
        ]
        
        found_count = 0
        for expected in expected_skills:
            if any(expected in skill_name for skill_name in skill_names):
                found_count += 1
        
        # Should find at least 70% of expected skills
        accuracy = found_count / len(expected_skills)
        assert accuracy >= 0.7, f"Skill extraction accuracy too low: {accuracy}"
    
    @pytest.mark.asyncio
    async def test_processing_performance(self, nlp_engine, sample_resume_text):
        """Test processing performance."""
        import time
        
        start_time = time.time()
        resume_data = await nlp_engine._process_resume_text(sample_resume_text)
        end_time = time.time()
        
        processing_time = end_time - start_time
        
        # Should process within reasonable time (adjust threshold as needed)
        assert processing_time < 30.0, f"Processing took too long: {processing_time}s"
        assert len(resume_data.skills) > 0
    
    # Test Cleanup
    def test_cleanup(self, nlp_engine):
        """Test resource cleanup."""
        # Should not raise any exceptions
        nlp_engine.cleanup()
    
    # Test Model Loading (Mock Tests)
    @patch('spacy.load')
    def test_spacy_model_loading(self, mock_spacy_load, nlp_engine):
        """Test spaCy model loading with fallback."""
        # Test successful loading
        mock_model = Mock()
        mock_spacy_load.return_value = mock_model
        
        model = nlp_engine.spacy_model
        assert model == mock_model
        
        # Test fallback to blank model
        mock_spacy_load.side_effect = OSError("Model not found")
        nlp_engine._spacy_model = None  # Reset cached model
        
        with patch('spacy.blank') as mock_blank:
            mock_blank_model = Mock()
            mock_blank.return_value = mock_blank_model
            
            model = nlp_engine.spacy_model
            assert model == mock_blank_model
            mock_blank.assert_called_with("en")
    
    @patch('sentence_transformers.SentenceTransformer')
    def test_sentence_transformer_loading(self, mock_st, nlp_engine):
        """Test sentence transformer model loading."""
        mock_model = Mock()
        mock_st.return_value = mock_model
        
        model = nlp_engine.sentence_transformer
        assert model == mock_model
        mock_st.assert_called_once()


# Integration Tests
class TestNLPEngineIntegration:
    """Integration tests for NLP Engine with real models."""
    
    @pytest.fixture(scope="class")
    def nlp_engine_real(self):
        """Create NLP engine with real models for integration testing."""
        with tempfile.TemporaryDirectory() as temp_dir:
            engine = NLPEngine(model_cache_dir=temp_dir)
            yield engine
            engine.cleanup()
    
    @pytest.mark.integration
    @pytest.mark.asyncio
    async def test_end_to_end_resume_processing(self, nlp_engine_real):
        """Test complete resume processing pipeline."""
        resume_text = """
        Jane Smith
        Senior Data Scientist
        jane.smith@email.com
        
        EXPERIENCE
        Senior Data Scientist at DataCorp (2021-2023)
        - Built machine learning models using Python, TensorFlow, and scikit-learn
        - Worked with large datasets using Pandas and NumPy
        - Deployed models to AWS using Docker and Kubernetes
        - Collaborated with cross-functional teams using Agile methodologies
        
        EDUCATION
        Master of Science in Data Science
        Stanford University (2019-2021)
        
        SKILLS
        Programming: Python, R, SQL
        ML Frameworks: TensorFlow, PyTorch, scikit-learn
        Cloud: AWS, Google Cloud Platform
        Tools: Docker, Kubernetes, Git, Jupyter
        """
        
        resume_data = await nlp_engine_real._process_resume_text(resume_text)
        
        # Verify comprehensive data extraction
        assert len(resume_data.skills) >= 10
        assert len(resume_data.experience) >= 1
        assert len(resume_data.education) >= 1
        
        # Verify skill categorization
        skill_categories = [skill.category for skill in resume_data.skills]
        assert SkillCategory.PROGRAMMING_LANGUAGES in skill_categories
        assert SkillCategory.FRAMEWORKS_LIBRARIES in skill_categories
        assert SkillCategory.CLOUD_PLATFORMS in skill_categories
        
        # Verify high-confidence skills
        high_conf_skills = resume_data.get_high_confidence_skills(threshold=0.6)
        assert len(high_conf_skills) > 0
    
    @pytest.mark.integration
    @pytest.mark.asyncio
    async def test_skill_similarity_accuracy(self, nlp_engine_real):
        """Test semantic similarity accuracy with real models."""
        # Similar skills should have high similarity
        similarity1 = await nlp_engine_real.calculate_similarity(
            "Python programming", "Python development"
        )
        assert similarity1 > 0.7
        
        # Related but different skills should have medium similarity
        similarity2 = await nlp_engine_real.calculate_similarity(
            "Python programming", "JavaScript development"
        )
        assert 0.3 < similarity2 < 0.8
        
        # Unrelated terms should have low similarity
        similarity3 = await nlp_engine_real.calculate_similarity(
            "Python programming", "cooking recipes"
        )
        assert similarity3 < 0.3


if __name__ == "__main__":
    pytest.main([__file__, "-v"])